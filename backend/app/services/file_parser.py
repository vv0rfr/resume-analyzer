"""
文件文本提取服务
支持 PDF (.pdf) 和 Word (.docx) 格式
"""
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
SUPPORTED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    # MIME 类型兼容
    "pdf": "pdf",
    "docx": "docx",
}


def extract_text(filename: str, content: bytes, mime_type: str = "") -> str:
    """
    从文件内容中提取文本

    Args:
        filename: 原始文件名
        content: 文件字节内容
        mime_type: MIME 类型（可选，用于确认格式）

    Returns:
        提取的文本内容

    Raises:
        ValueError: 不支持的文件格式、文件过大、解析失败
    """
    # 检查文件大小
    if len(content) > MAX_FILE_SIZE:
        raise ValueError(f"文件过大（最大 10MB），当前文件 {len(content) / 1024 / 1024:.1f}MB")

    # 确定文件类型
    ext = _get_extension(filename)
    mime_ext = SUPPORTED_TYPES.get(mime_type, "")
    file_type = mime_ext or ext or "unknown"

    if file_type not in ("pdf", "docx"):
        raise ValueError(
            f"不支持的文件格式「.{ext}」，请上传 PDF 或 Word (.docx) 文件"
        )

    # 提取文本
    try:
        if file_type == "pdf":
            return _extract_pdf(content)
        else:
            return _extract_docx(content)
    except Exception as e:
        logger.error(f"文件解析失败 [{filename}]: {e}")
        raise ValueError(f"无法读取该文件，请确认文件未损坏或不是空文件") from e


def _get_extension(filename: str) -> Optional[str]:
    """获取小写文件扩展名"""
    if "." in filename:
        return filename.rsplit(".", 1)[-1].lower()
    return None


def _extract_pdf(content: bytes) -> str:
    """使用 pdfplumber 提取 PDF 文本"""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    result = "\n".join(text_parts).strip()
    if not result:
        raise ValueError("PDF 文件中未检测到文本内容，可能为扫描件或图片型 PDF")

    logger.info(f"PDF 提取完成: {len(pdf.pages)} 页, {len(result)} 字符")
    return result


def _extract_docx(content: bytes) -> str:
    """使用 python-docx 提取 Word 文本"""
    from docx import Document

    doc = Document(io.BytesIO(content))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)

    result = "\n".join(text_parts).strip()
    if not result:
        raise ValueError("Word 文件中未检测到文本内容")

    logger.info(f"DOCX 提取完成: {len(doc.paragraphs)} 段落, {len(result)} 字符")
    return result
